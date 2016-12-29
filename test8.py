from docx import Document
import time
import os

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

#os part
dest_dir = input('请输入外部审校文件所在路径。\n>').replace("\\", "/")
start = time.time()
for root, dirs, files in os.walk(dest_dir):
    pass

for name in files:
    t = Document(os.path.join(dest_dir, name)).tables[0]
    i = 0
    for cell in t.columns[1].cells:
        if '100%' in cell.text or 'CM' in cell.text:
            delete_paragraph(t.cell(i, 2).paragraphs[0])
        i += 1
print (time.time() - start)