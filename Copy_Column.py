import os
from openpyxl import Workbook
from openpyxl.styles import PatternFill
from docx import Document
import re
import time

def del_blank(text):
    return text != ''

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def find_word(name):
    return '.doc' in name

#excel part
wb = Workbook()
ws = wb.active
fil = PatternFill(start_color='FFFF00', end_color='FFFF00',fill_type='solid')

#os part
dest_dir = input('请输入外部审校文件所在路径。\n>').replace("\\", "/")
start = time.time()
for root, dirs, files in os.walk(dest_dir):
    pass
files = filter(find_word,files)
for name in files:
    print ('正在处理' + name + '...')
    t = Document(os.path.join(dest_dir, name)).tables[0]
    j = []
    j.clear()
    ws.append({'B': name})
    ws['B' + str(ws.max_row)].fill = fil
    i = 0
    for cell in t.columns[1].cells:
        if '100%' in cell.text or 'CM' in cell.text or 'Draft' in cell.text:
            delete_paragraph(t.cell(i, 2).paragraphs[0])
        i += 1
    for cell in t.columns[2].cells:
        j.append(re.sub('<'r'/?[a-z]{0,3}[0-9]{0,6}/?''>', '', cell.text))
    r1 = list(filter(del_blank, j))
    r2 = [n for n in range(len(t.rows))]
    for row in zip(r2, r1):
        ws.append(row)

wb.save(dest_dir[:dest_dir.rfind('/',0,dest_dir.find('External Review')-1)+ 1] + '删重文件.xlsx')
print (time.time() - start)